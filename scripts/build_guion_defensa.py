#!/usr/bin/env python3
"""Reconstruye la sección 'DEFENSA' del guion, alineada al deck real (19 slides).

El guion de defensa venía de una maqueta antigua (con índice, separadores y
slide de referencias que el deck actual no tiene, y sin los slides de
'Progresión Iter' ni 'Lenguaje natural'). Aquí se reescribe 1:1 con el deck.

- Deja INTACTA la sección divulgativa (se inserta antes de su Heading 1).
- Idempotente: re-ejecutar produce el mismo resultado.

Uso:  ../.venv_thesis/bin/python scripts/build_guion_defensa.py
"""
from __future__ import annotations

from pathlib import Path
from docx import Document

REPO = Path(__file__).resolve().parents[1]
DOC = REPO / "docs/Guion_Completo_Presentaciones.docx"

H1 = "DEFENSA (plantilla UNIR) — ~20 min"
DIV_H1_MATCH = "divulgativa"

SLIDES = [
    ("Slide 1 — Estimación de Pose 6-DoF con Transformers y Modelos de Difusión para Bin Picking Robótico", [
        ("cue", "[PORTADA · ~45 s]"),
        ("say", "«Buenos días. Agradezco al tribunal su tiempo. Soy Giocrisrai Godoy "
                "y defiendo mi Trabajo Fin de Máster: \"Estimación de pose 6-DoF con "
                "Transformers y modelos de difusión para bin picking robótico\". En "
                "los próximos veinte minutos quiero convencerles de una idea sencilla "
                "pero potente: que unir dos avances de la IA que hasta ahora vivían "
                "separados —los Transformers para percibir y los modelos de difusión "
                "para planificar— permite atacar un problema que la industria "
                "persigue desde hace décadas. Seguiré cuatro partes: el problema y "
                "los objetivos; el método y su fundamento matemático; los resultados "
                "—con total honestidad, incluido lo que no salió a la primera—; y las "
                "conclusiones.»"),
        ("note", "[Respira. Mira al tribunal, no a la pantalla. → \"Empecemos por ese "
                 "problema.\"]"),
    ]),
    ("Slide 2 — El problema: bin-picking 6-DoF", [
        ("cue", "[EL PROBLEMA · ~60 s]"),
        ("say", "«El bin picking consiste en algo aparentemente trivial: coger una "
                "pieza de un montón desordenado, agarrarla y depositarla. Pero "
                "esconde la paradoja de Moravec: lo que a un humano le resulta fácil "
                "—ver y manipular— es lo más difícil de replicar en una máquina. Para "
                "situarlo: ganarle al campeón mundial de ajedrez se resolvió en 1997; "
                "que un robot recoja con fiabilidad una pieza de un cajón sigue siendo "
                "investigación abierta. Y no es un problema de nicho: cada paquete de "
                "e-commerce pasó por unas manos haciendo exactamente esto.»"),
        ("note", "[→ \"Para resolverlo, el robot debe responder dos preguntas: dónde "
                 "está la pieza y cómo agarrarla.\"]"),
    ]),
    ("Slide 3 — Objetivo e hipótesis", [
        ("cue", "[OBJETIVO E HIPÓTESIS · ~75 s]"),
        ("say", "«Mi objetivo fue integrar esas dos respuestas —percepción y "
                "planificación— en un único pipeline, validado de principio a fin. Y "
                "me comprometí con tres hipótesis falsables, con números, no con "
                "adjetivos. H1, de precisión: mejorar la métrica oficial en al menos "
                "tres puntos sobre el referente. H2, de planificación: generar "
                "agarres multimodales, rápidos y de calidad. Y H3, de viabilidad: que "
                "todo corra sin GPU dedicada, por debajo de diez segundos por ciclo. "
                "Les adelanto el final: dos de las tres se cumplen, y una se cumple "
                "solo en parte. Y esa parte, lejos de ser un problema, es quizá lo "
                "más interesante que voy a contarles.»"),
        ("note", "[Di \"solo en parte\" con calma y seguridad, no como disculpa: es "
                 "madurez científica.]"),
    ]),
    ("Slide 4 — Estado del arte y brecha", [
        ("cue", "[ESTADO DEL ARTE Y BRECHA · ~70 s]"),
        ("say", "«¿Dónde está el campo? La estimación de pose pasó de los métodos "
                "clásicos a los Transformers, con FoundationPose marcando el estado "
                "del arte en 2024. La planificación de agarre encontró en los modelos "
                "de difusión una forma natural de capturar que hay muchas maneras "
                "válidas de coger un objeto. Son dos líneas muy potentes… que nadie "
                "había unido con un marco matemático común para bin picking. Esa es "
                "exactamente mi brecha, y mi contribución. No propongo una pieza nueva "
                "aislada: propongo el puente formal entre dos mundos.»"),
        ("note", "[Margen: nombra 1-2 trabajos del estado del arte si el tribunal es "
                 "del área.]"),
    ]),
    ("Slide 5 — Pipeline integrado", [
        ("cue", "[PIPELINE INTEGRADO · ~70 s]"),
        ("say", "«Así funciona, de izquierda a derecha. La cámara RGB-D entra a "
                "FoundationPose, que estima la pose del objeto en SE(3). Esa pose "
                "condiciona a la Diffusion Policy, que genera una trayectoria de "
                "dieciséis pasos. Y cada paso se ejecuta mediante cinemática inversa "
                "en el simulador. Tres familias de técnicas —percepción, generación y "
                "control— en un solo flujo coherente. Subrayo que lo novedoso no es "
                "ninguno de los bloques por separado, sino su integración.»"),
        ("note", "[Recorre el diagrama con el puntero, izq→der, sin prisa. Citas: "
                 "FoundationPose = Wen et al., CVPR 2024; Diffusion Policy = Chi et "
                 "al., RSS 2023.]"),
    ]),
    ("Slide 6 — Fundamento matemático", [
        ("cue", "[FUNDAMENTO MATEMÁTICO · ~110 s — TU TERRENO]"),
        ("say", "«Y aquí está el corazón formal, el que justifica que esto sea un "
                "trabajo de Ingeniería Matemática. Las rotaciones viven en SO(3); uso "
                "la representación 6D continua, que evita las discontinuidades del "
                "cuaternión y el gimbal lock. La pose completa es un elemento de "
                "SE(3), un grupo de Lie. La percepción se apoya en la atención scaled "
                "dot-product. Y la planificación es score matching más dinámica de "
                "Langevin: una ecuación diferencial estocástica que, partiendo de "
                "ruido, recupera una trayectoria válida. El puente entre percepción y "
                "acción es geométrico: la pose en SE(3) condiciona la SDE que genera "
                "el movimiento. Esa conexión, derivada formalmente, es la aportación "
                "original.»"),
        ("note", "[Habla DESPACIO; es donde más puntúas. Si te preguntan aquí, "
                 "alégrate. Citas: rep. 6D = Zhou et al. 2019; <5D imposible = "
                 "Stuelpnagel 1964; SE(3)/Lie = Solà et al. 2018; difusión = Ho 2020 "
                 "/ Song 2021.]"),
    ]),
    ("Slide 7 — Metodología experimental", [
        ("cue", "[METODOLOGÍA · ~60 s]"),
        ("say", "«Antes de los números, cómo los obtuve, porque el rigor es parte del "
                "aporte. Evalué sobre dos datasets estándar: YCB-Video y T-LESS; este "
                "último, piezas industriales sin textura, es el caso difícil y el más "
                "relevante para bin picking. Usé las métricas oficiales del BOP "
                "Challenge, que tratan correctamente las simetrías, con los splits "
                "oficiales, varias semillas, e intervalos de confianza al 95 % por "
                "bootstrap. No reporto ningún número suelto: todo va con su "
                "incertidumbre.»"),
    ]),
    ("Slide 8 — H1 · Precisión de pose — ACEPTADA", [
        ("cue", "[H1 · PRECISIÓN — ACEPTADA · ~75 s]"),
        ("say", "«Primera hipótesis, precisión. Atención al matiz de métrica, porque "
                "es importante: en la figura ven la comparación sobre Mean AR, la "
                "métrica oficial del leaderboard, donde FoundationPose supera a "
                "GDR-Net++ en 3,0 puntos en YCB-Video y 3,6 en T-LESS, por encima del "
                "umbral que me fijé. Adicionalmente, mi validación local recomputa un "
                "AUC ADD-S de 0,908 y 0,957, una métrica relacionada pero distinta "
                "—lo declaro como limitación de constructo—. Sobre Mean AR, H1 se "
                "acepta. Y, de paso, reproduje el estado del arte en mi propio "
                "hardware, lo cual ya es un resultado.»"),
        ("note", "[Pausa tras los números. No mezcles las dos métricas: di "
                 "explícitamente \"Mean AR\" para el +3 pp.]"),
    ]),
    ("Slide 9 — H2 · Planificación multimodal — PARCIALMENTE ACEPTADA", [
        ("cue", "[H2 · MULTIMODAL — PARCIAL · ~95 s — EL MOMENTO HONESTO]"),
        ("say", "«Segunda hipótesis. Aquí soy transparente. La calidad de las "
                "trayectorias cumple, con un score de 0,96. Pero el muestreo tarda "
                "118 milisegundos, por encima de los 50 que pedí, y detecto dos modos "
                "en vez de los tres exigidos. Por eso H2 se acepta solo en parte. ¿La "
                "causa? La diagnostiqué: la política aprendió imitando un planificador "
                "determinista, no demostraciones reales, así que copia bien pero "
                "explora poco. Y al sustituir el condicionamiento por un encoder "
                "visual, el agarre ejecutado subió del 36 % al 78 %, lo que confirma "
                "que ahí estaba el cuello de botella. La ciencia honesta no es la que "
                "nunca falla; es la que entiende por qué falla.»"),
        ("note", "[Mira a los ojos al tribunal. No te disculpes: explica. Este slide "
                 "te gana puntos.]"),
    ]),
    ("Slide 10 — H3 · Viabilidad sin GPU dedicada — ACEPTADA", [
        ("cue", "[H3 · VIABILIDAD — ACEPTADA · ~55 s]"),
        ("say", "«Tercera hipótesis, viabilidad sin GPU dedicada. El ciclo completo "
                "se ejecuta con un percentil 95 de 6,3 segundos en YCB-Video y 6,7 en "
                "T-LESS, holgadamente por debajo de los diez segundos. El profiling "
                "muestra que el cuello de botella es la percepción, que ocupa el 80 % "
                "del tiempo; la difusión es apenas el 2 %. Y todo corre sobre una "
                "laptop más una GPU de alquiler en la nube. H3 se acepta.»"),
    ]),
    ("Slide 11 — Progresión de la planificación: Iter 5 → 7c", [
        ("cue", "[FRACASO → REMONTADA · ~80 s — EL CLÍMAX]"),
        ("say", "«Esta curva resume la parte que casi nadie cuenta en una defensa: un "
                "fracaso, y la remontada. Mi primer sistema completo lograba un 60 % "
                "de ciclos exitosos. Quise mejorarlo con aprendizaje por refuerzo… y "
                "bajó al 28 %. Había optimizado el depósito y, al hacerlo, olvidó "
                "cómo agarrar. Se llama olvido catastrófico y le ocurre a los mejores "
                "laboratorios del mundo. No me rendí: lo diagnostiqué, apliqué un "
                "currículo —primero agarrar, después depositar—, generé ocho "
                "candidatas y ejecuté la mejor, y corregí un error de geometría. "
                "Resultado: 84 %, por encima del sistema original. El mensaje es que "
                "en IA real no se avanza por arte de magia, sino diagnosticando.»"),
        ("note", "[Cuenta esto como una historia, con ritmo. Es tu mejor momento "
                 "narrativo; apóyate en la curva del slide.]"),
    ]),
    ("Slide 12 — Demo: el pipeline en acción", [
        ("cue", "[DEMO (VÍDEO) · ~60 s]"),
        ("say", "«Y para que no se queden solo con gráficos, véanlo funcionar.»"),
        ("note", "[Reproducir el vídeo con clic. NO hables encima; deja que respire. "
                 "Solo señala las fases: \"ahí estima la pose… ahí genera la "
                 "trayectoria… ahí agarra… y deposita\". Fíjate en el panel de "
                 "telemetría: el ciclo total y el veredicto H3 en vivo.]"),
        ("note", "[Plan B: si no reproduce, descríbelo con calma y sigue, sin "
                 "dramatismo. Vídeos sueltos en docs/entrega3/videos_proyeccion/.]"),
    ]),
    ("Slide 13 — Bin picking guiado por lenguaje natural", [
        ("cue", "[LENGUAJE NATURAL (VÍDEO) · ~60 s — EXTENSIÓN QUE ABRE EL FUTURO]"),
        ("say", "«Cierro los resultados con una extensión que muestra hacia dónde "
                "crece el trabajo. Sobre el mismo pipeline, y sin tocar su núcleo, "
                "añadí una capa opcional de lenguaje natural: le digo \"dame el cubo "
                "rojo de la izquierda\" y el sistema interpreta la orden, elige la "
                "pieza correcta entre varias y la agarra. En el vídeo lo ven con cinco "
                "instrucciones, en español e inglés, distinguiendo color, forma y "
                "posición. Es código abierto, corre en local, y está validado de "
                "extremo a extremo en el simulador con una selección honesta: si la "
                "pieza pedida no está, no la inventa. Es la prueba de que la "
                "arquitectura no solo funciona, sino que es extensible.»"),
        ("note", "[Reproducir el vídeo embebido (39 s). Señala la orden y la pieza "
                 "elegida. Si preguntan por el LLM: parser determinista + modelo "
                 "local opcional, con fallback; no depende de la nube.]"),
    ]),
    ("Slide 14 — Honestidad y limitaciones", [
        ("cue", "[HONESTIDAD Y LÍMITES · ~55 s]"),
        ("say", "«Quiero declarar también lo que NO hice, porque conocer los límites "
                "es parte del resultado. El agarre es snap+attach: valida la cadena "
                "percepción-planificación-ejecución, pero no la mecánica de fricción "
                "del contacto. Todo es simulación, sin brazo robótico real. Y "
                "FoundationPose tiene licencia no comercial, que restringe su uso "
                "industrial directo, aunque es sustituible por un estimador de "
                "licencia permisiva sin tocar el resto del pipeline.»"),
    ]),
    ("Slide 15 — Aportes 1 y 2", [
        ("cue", "[APORTES 1-2 · ~45 s]"),
        ("say", "«Sintetizo el aporte en cuatro contribuciones; las dos primeras, "
                "aquí. La primera, metodológica y central: la integración matemática "
                "formal de percepción con Transformers y planificación por difusión "
                "sobre SE(3) —el puente entre los dos mundos—. La segunda, de "
                "reproducibilidad: recompongo el estado del arte en hardware "
                "accesible, con tests, integración continua y un paquete propio "
                "publicado en PyPI, instalable con una sola línea.»"),
    ]),
    ("Slide 16 — Aportes 3 y 4", [
        ("cue", "[APORTES 3-4 · ~40 s]"),
        ("say", "«Y las otras dos. La tercera, de rigor: intervalos de confianza al "
                "95 %, ablaciones y un análisis explícito de amenazas a la validez, "
                "por encima del estándar del área. Y la cuarta, la validación de "
                "principio a fin, con evidencia visual reproducible y repositorio "
                "público —incluida la extensión de lenguaje natural que acaban de "
                "ver—.»"),
        ("note", "[Margen: enfatiza la que más valoras.]"),
    ]),
    ("Slide 17 — Valor y potencialidad", [
        ("cue", "[VALOR Y POTENCIALIDAD · ~50 s]"),
        ("say", "«¿Y por qué importa, y hasta dónde puede llegar? El mercado de bin "
                "picking crece hacia los siete mil millones de dólares, y las "
                "soluciones cuestan entre quince y ciento cincuenta mil. Este pipeline "
                "corre por menos de dos mil, como coste de I+D y reproducción, no de "
                "despliegue. El mensaje no es que reemplace una celda industrial, sino "
                "que investigar y reproducir el estado del arte en manipulación —y "
                "extenderlo, como con el lenguaje natural— ya está al alcance de una "
                "universidad o una pyme.»"),
        ("note", "[Aclara \"coste de I+D, no de despliegue\" para que no te lo "
                 "refuten.]"),
    ]),
    ("Slide 18 — Trabajo futuro", [
        ("cue", "[TRABAJO FUTURO · ~45 s]"),
        ("say", "«Tres líneas claras hacia delante: validar el agarre físico real "
                "sustituyendo el snap+attach; acelerar la difusión con flow-matching, "
                "de veinticinco pasos a uno o dos; y transferir el sistema a un robot "
                "real. Y subrayo que es la misma receta —percepción con Transformers "
                "más acción por difusión— que mueve a los robots humanoides de hoy.»"),
    ]),
    ("Slide 19 — Cierre", [
        ("cue", "[CIERRE · ~30 s]"),
        ("say", "«Para terminar: he presentado un pipeline que une matemáticamente "
                "percepción y planificación, validado con rigor y con honestidad "
                "—dos hipótesis aceptadas y una en parte—, reproducible por menos de "
                "lo que cuesta una moto, extensible a lenguaje natural, y con todo el "
                "código y los datos públicos. Muchas gracias por su atención; quedo a "
                "su disposición para las preguntas.»"),
        ("note", "[Última frase más lenta. Sonríe. Las 42 referencias y el material "
                 "de respaldo están en la memoria por si el tribunal pide fuentes.]"),
    ]),
]


def _div_h1(doc):
    for p in doc.paragraphs:
        if p.style and p.style.name == "Heading 1" and DIV_H1_MATCH in p.text.lower():
            return p
    return None


def _del_defensa(doc):
    """Borra desde el Heading 1 de DEFENSA hasta (sin incluir) la H1 divulgativa."""
    paras = doc.paragraphs
    start = end = None
    for i, p in enumerate(paras):
        if p.style and p.style.name == "Heading 1":
            if "DEFENSA" in p.text and start is None:
                start = i
            elif DIV_H1_MATCH in p.text.lower():
                end = i
                break
    if start is None:
        return False
    stop = end if end is not None else len(paras)
    for p in paras[start:stop]:
        p._p.getparent().remove(p._p)
    return True


def main() -> int:
    from docx.text.paragraph import Paragraph

    doc = Document(DOC)
    existed = _del_defensa(doc)
    print(f"  sección defensa previa {'eliminada' if existed else 'no encontrada'}")

    anchor = _div_h1(doc)  # insertamos ANTES de la divulgativa
    if anchor is None:
        raise SystemExit("No se encontró la sección divulgativa como ancla.")

    def add_before(text, style):
        p = doc.add_paragraph(text, style=style)  # se crea al final…
        anchor._p.addprevious(p._p)               # …y se mueve antes del ancla
        return p

    add_before(H1, "Heading 1")
    for title, blocks in SLIDES:
        add_before(title, "Heading 2")
        for _kind, text in blocks:
            add_before(text, "Normal")

    doc.save(DOC)
    print(f"  defensa reescrita: {len(SLIDES)} slides alineados al deck (19)")
    print(f"  -> {DOC.relative_to(REPO)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

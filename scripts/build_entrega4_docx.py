#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Genera el .docx de la Entrega 4 del TFM.

Parte de una copia de la Entrega 3 e inserta, en tres puntos de anclaje, la
contribucion de la Entrega 4: una capa de bin picking guiado por lenguaje
natural (arquitectura, ejecucion E2E, resultados, tablas, figura y limitaciones),
una comparacion con el estado del arte y un cierre en conclusiones.

El script es idempotente: siempre re-copia desde la Entrega 3 al iniciar, de modo
que volver a ejecutarlo regenera limpiamente el documento sin duplicar contenido.

IMPORTANTE: requiere python-docx, que SOLO esta instalado en `../.venv_thesis`.
Ejecutar con:
    ../.venv_thesis/bin/python scripts/build_entrega4_docx.py
"""

import os
import shutil
import sys

from docx import Document
from docx.shared import Inches, Pt  # noqa: F401  (Pt disponible por si se ajustan tamanos)
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Rutas (relativas a la raiz del repo repo_tfm) ---------------------------
REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
E3_PATH = os.path.join(REPO_ROOT, "docs", "entrega3", "TFM_Entrega3_UNIR.docx")
E4_DIR = os.path.join(REPO_ROOT, "docs", "entrega4")
E4_PATH = os.path.join(E4_DIR, "TFM_Entrega4_UNIR.docx")
FIG_PATH = os.path.join(E4_DIR, "assets", "fig_language_pick.png")


# --- Helpers -----------------------------------------------------------------
def find_anchor(doc, style_name, text):
    """Devuelve el primer parrafo con el estilo y texto dados, o falla ruidosamente.

    El emparejamiento es por `style.name` exacto y por `text.strip()` exacto, de
    modo que un cambio futuro en el documento haga fallar el script de forma clara
    en lugar de insertar contenido en un lugar incorrecto.
    """
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == style_name and p.text.strip() == text:
            return p
    raise RuntimeError(
        "No se encontro el ancla: estilo=%r texto=%r. "
        "El documento base ha cambiado; revise los anclajes." % (style_name, text)
    )


def insert_para(anchor, text, style=None):
    """Inserta un parrafo (con estilo opcional) inmediatamente antes del ancla."""
    return anchor.insert_paragraph_before(text, style=style)


def insert_caption(anchor, text):
    """Inserta un parrafo de leyenda en cursiva antes del ancla.

    Usa el estilo 'Caption' si existe; en caso contrario, un parrafo normal con
    la corrida en cursiva.
    """
    try:
        cap = anchor.insert_paragraph_before(text, style="Caption")
        # Asegura cursiva tambien sobre el estilo Caption.
        for run in cap.runs:
            run.italic = True
        if not cap.runs:
            cap.add_run(text).italic = True
        return cap
    except (KeyError, ValueError):
        cap = anchor.insert_paragraph_before("")
        cap.add_run(text).italic = True
        return cap


def insert_table_before(doc, anchor, header, rows, table_style):
    """Crea una tabla (cabecera + filas) y la mueve a la posicion previa al ancla."""
    n_cols = len(header)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    if table_style is not None:
        t.style = table_style
    # Cabecera
    for c, txt in enumerate(header):
        cell = t.cell(0, c)
        cell.text = txt
        # Cabecera en negrita.
        for run in cell.paragraphs[0].runs:
            run.bold = True
    # Filas de datos
    for r, row in enumerate(rows, start=1):
        for c, txt in enumerate(row):
            t.cell(r, c).text = txt
    # Mover la tabla (creada al final del documento) a su posicion correcta.
    anchor._p.addprevious(t._tbl)
    return t


def insert_figure(anchor, image_path, width):
    """Inserta una imagen centrada inmediatamente antes del ancla."""
    p = anchor.insert_paragraph_before("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(image_path, width=width)
    return p


# --- Contenido a insertar (verbatim) -----------------------------------------
A4_ARCH_TITLE = "Bin picking guiado por lenguaje natural: arquitectura"
A4_ARCH_BODY = (
    "El sistema de la Entrega 4 incorpora una capa de comprensión de lenguaje "
    "natural que se sitúa por encima del pipeline de bin picking ya descrito. Con "
    "ella, el operador puede expresar una orden en términos cotidianos —«dame el "
    "cubo rojo de la izquierda»— y el robot identifica y manipula la pieza "
    "aludida. Esta capa se apoya en dos componentes con responsabilidades bien "
    "diferenciadas. El primero, el analizador de instrucciones, convierte la "
    "frase en una representación estructurada que recoge la intención, los "
    "atributos del objetivo (color, forma y tamaño) y, cuando procede, la "
    "relación espacial. Su versión por defecto es determinista y se sustenta en "
    "un léxico controlado español-inglés que coteja palabras completas; junto a "
    "ella conviven, de forma intercambiable, un analizador basado en un modelo de "
    "lenguaje local y otro basado en una API, que recurren al determinista si "
    "fallan. El segundo componente, el anclaje, vincula la instrucción con los "
    "objetos presentes en la escena: puntúa la coincidencia de atributos y, ante "
    "el empate, decide según la posición relativa de los centroides. La "
    "integración es deliberadamente poco invasiva —se habilita mediante una "
    "opción de configuración— y, cuando hay instrucción, restringe los candidatos "
    "a la pieza descrita antes de planificar el agarre."
)
A4_E2E_TITLE = "Ejecución extremo a extremo en simulación"
A4_E2E_BODY = (
    "Para comprobar el comportamiento en ejecución se desarrolló una rutina "
    "extremo a extremo sobre CoppeliaSim. Esta rutina genera una escena con "
    "varios objetos de formas distintas —cubos, esferas y cilindros creados como "
    "primitivas—, ancla la instrucción a la pieza que le corresponde y completa "
    "el ciclo de pick-and-place sobre ella, apoyándose en el mismo controlador "
    "cinemático y la misma técnica de agarre del pipeline base. El criterio de "
    "acierto es exigente: una selección solo cuenta como correcta si el objeto "
    "manipulado satisface todos los atributos enunciados en la orden, de modo que "
    "las coincidencias parciales no inflan los resultados."
)
A4_RES_TITLE = "Resultados y validación de la capa de lenguaje"
A4_RES_BODY = (
    "La validación se abordó en dos planos: uno reproducible sin simulador y otro "
    "de ejecución real en CoppeliaSim. La tabla de capacidades sintetiza lo que "
    "la Entrega 4 añade frente a la anterior, mientras que la de validación reúne "
    "la evidencia. En el banco reproducible —noventa escenas con formas "
    "variadas— la selección fue correcta en el 100 % de los casos para las tres "
    "familias consideradas, color, forma y relación espacial. Cabe matizar el "
    "alcance de esa cifra: confirma que la cadena de análisis, anclaje y "
    "selección funciona de principio a fin sobre un banco donde la pieza descrita "
    "es inequívoca, mientras que la solidez ante lenguaje ambiguo o no visto "
    "descansa en las exploraciones exp16 a exp26 sobre el modelo aprendido. En el "
    "plano real, la prueba de integración cerró el ciclo en noventa y ocho "
    "segundos; una ejecución representativa de la orden «dame el cubo rojo de la "
    "izquierda» escogió la pieza adecuada con una distancia pinza-objeto de "
    "cuatro milímetros —holgadamente por debajo del umbral de cinco "
    "centímetros— y con la cinemática inversa convergiendo; y el banco en "
    "simulador, de nueve escenas, volvió a alcanzar el 100 %."
)

CAP_HEADER = ["Capacidad", "Entrega 3", "Entrega 4"]
CAP_ROWS = [
    ["Pose 6-DoF, agarre y servoing", "Sí (H1/H2/H3 validadas)", "Sí (preservado)"],
    ["Selección por lenguaje natural", "No", "Sí"],
    ["Anclaje por atributos y relación espacial", "No", "Sí"],
    ["Variedad de formas en simulación", "Cubos", "Cubo, esfera y cilindro"],
    ["Ejecución E2E por instrucción", "No", "Sí"],
    ["Métrica de selección honesta", "—", "Sí"],
]
CAP_CAPTION = "Tabla. Comparación de capacidades entre la Entrega 3 y la Entrega 4."

VAL_HEADER = ["Prueba", "Resultado", "Fuente"]
VAL_ROWS = [
    ["Batería de selección (pura, 90 escenas)", "100 % (color, forma y espacial)", "report_pure_n90.json"],
    ["Test de integración E2E", "Completado en 98 s", "test de integración"],
    ["Ejecución E2E representativa", "Selección correcta; agarre a 4 mm; IK convergente", "run en vivo"],
    ["Batería de selección en simulador (9 escenas)", "100 %", "report_sim_n9.json"],
    ["Reel de demostración", "174 s, 5 instrucciones", "language_reel.mp4"],
]
VAL_CAPTION = "Tabla. Evidencia de validación de la capa de lenguaje (junio de 2026)."

FIG_CAPTION = (
    "Figura. Ejecución extremo a extremo de la instrucción «dame el cubo rojo de "
    "la izquierda» en una escena con variedad de formas."
)

A4_LIM_TITLE = "Limitaciones y trabajo futuro de la capa de lenguaje"
A4_LIM_BODY = (
    "Conviene reconocer con franqueza las limitaciones de esta capa. El agarre es "
    "cinemático, por lo que la medida honesta de su calidad es la cercanía entre "
    "pinza y objeto y la convergencia de la cinemática inversa, no el "
    "desplazamiento final de la pieza. El banco de selección, aunque valida la "
    "cadena completa, es controlado y no contempla casos realmente ambiguos —dos "
    "objetos idénticos en color y forma—, algo que se reserva como línea futura. "
    "La estimación del tamaño mediante el modelo de visión-lenguaje resulta la "
    "menos fiable, al faltar una referencia de escala. Y el backend de API queda "
    "como punto de extensión, todavía sin un proveedor concreto."
)

A2_BODY = (
    "Frente a propuestas de manipulación guiada por lenguaje como CLIPort, "
    "VoxPoser, SayCan u OWL-ViT y CLIP-Fields, el enfoque aquí planteado se "
    "distingue por ser íntegramente open-license y por ejecutarse en un portátil "
    "sin GPU dedicada, además de ofrecer un anclaje interpretable y una "
    "comprensión apoyada en modelos de lenguaje intercambiables que no obliga a "
    "reentrenar al cambiar el vocabulario. No se aporta una única cifra de "
    "precisión directamente comparable con esos trabajos, pues se evalúan sobre "
    "bancos heterogéneos; la comparación, por tanto, se mantiene cualitativa en "
    "esos ejes."
)

A5_BODY = (
    "Sobre el pipeline ya validado, la Entrega 4 suma una capa de bin picking "
    "guiado por lenguaje natural que permite señalar y manipular la pieza "
    "descrita en una orden, comprobada de extremo a extremo en simulación. De "
    "cara al futuro se contempla un banco de selección con casos genuinamente "
    "ambiguos, el anclaje por imagen sobre la cámara real del simulador y la "
    "incorporación de un proveedor concreto en el backend de API."
)

NEW_H3_TITLES = [A4_ARCH_TITLE, A4_E2E_TITLE, A4_RES_TITLE, A4_LIM_TITLE]


# --- Construccion ------------------------------------------------------------
def build():
    if not os.path.exists(E3_PATH):
        raise FileNotFoundError("No existe la Entrega 3: %s" % E3_PATH)
    if not os.path.exists(FIG_PATH):
        raise FileNotFoundError("No existe la figura: %s" % FIG_PATH)
    os.makedirs(E4_DIR, exist_ok=True)

    # Idempotencia: re-copiar siempre desde E3.
    shutil.copy(E3_PATH, E4_PATH)

    doc = Document(E4_PATH)

    # Estilo de tabla reutilizado del propio documento.
    table_style = doc.tables[0].style if doc.tables else None

    # ---- Antes de A4 (Capitulo 4), en orden de lectura ----
    a4 = find_anchor(doc, "Heading 3", "Resumen consolidado de la evidencia experimental")
    insert_para(a4, A4_ARCH_TITLE, style="Heading 3")
    insert_para(a4, A4_ARCH_BODY, style="Normal")
    insert_para(a4, A4_E2E_TITLE, style="Heading 3")
    insert_para(a4, A4_E2E_BODY, style="Normal")
    insert_para(a4, A4_RES_TITLE, style="Heading 3")
    insert_para(a4, A4_RES_BODY, style="Normal")
    insert_table_before(doc, a4, CAP_HEADER, CAP_ROWS, table_style)
    insert_caption(a4, CAP_CAPTION)
    insert_table_before(doc, a4, VAL_HEADER, VAL_ROWS, table_style)
    insert_caption(a4, VAL_CAPTION)
    insert_figure(a4, FIG_PATH, Inches(5.5))
    insert_caption(a4, FIG_CAPTION)
    insert_para(a4, A4_LIM_TITLE, style="Heading 3")
    insert_para(a4, A4_LIM_BODY, style="Normal")

    # ---- Antes de A2 (Capitulo 2, estado del arte) ----
    a2 = find_anchor(doc, "Heading 2", "Conclusiones del estado del arte y brecha de investigación")
    insert_para(a2, A2_BODY, style="Normal")

    # ---- Antes de A5 (Capitulo 5, conclusiones) ----
    a5 = find_anchor(doc, "Heading 2", "Discusión y cierre de hipótesis")
    insert_para(a5, A5_BODY, style="Normal")

    doc.save(E4_PATH)
    return doc


# --- Verificacion ------------------------------------------------------------
def heading_texts(doc):
    out = []
    for p in doc.paragraphs:
        if p.style is not None and p.style.name in ("Heading 1", "Heading 2", "Heading 3"):
            out.append(p.text.strip())
    return out


def verify():
    e3 = Document(E3_PATH)
    e3_tables = len(e3.tables)
    e3_shapes = len(e3.inline_shapes)
    e3_headings = set(t for t in heading_texts(e3) if t)

    e4 = Document(E4_PATH)  # se reabre: si no abre, lanza excepcion
    e4_tables = len(e4.tables)
    e4_shapes = len(e4.inline_shapes)
    e4_heading_list = heading_texts(e4)
    e4_headings = set(t for t in e4_heading_list if t)

    print("== VERIFICACION ENTREGA 4 ==")
    print("Documento E4 abierto correctamente: OK")

    # 4 nuevos titulos Heading 3
    for title in NEW_H3_TITLES:
        present = title in e4_heading_list
        print("Heading 3 presente [%s]: %s" % ("OK" if present else "FALTA", title))
        assert present, "Falta el Heading 3: %s" % title

    # Tablas: E4 == E3 + 2
    print("Tablas E3=%d  E4=%d  (esperado E3+2=%d)" % (e3_tables, e4_tables, e3_tables + 2))
    assert e4_tables == e3_tables + 2, "Conteo de tablas inesperado"

    # Imagenes: E4 == E3 + 1
    print("Imagenes (inline_shapes) E3=%d  E4=%d  (esperado E3+1=%d)" % (e3_shapes, e4_shapes, e3_shapes + 1))
    assert e4_shapes == e3_shapes + 1, "Conteo de imagenes inesperado"

    # Parrafos clave SOTA y conclusiones
    full_text = "\n".join(p.text for p in e4.paragraphs)
    sota_ok = "íntegramente open-license" in full_text
    concl_ok = "comprobada de extremo a extremo en simulación" in full_text
    print("Parrafo estado del arte ('íntegramente open-license'): %s" % ("OK" if sota_ok else "FALTA"))
    print("Parrafo conclusiones ('comprobada de extremo a extremo en simulación'): %s" % ("OK" if concl_ok else "FALTA"))
    assert sota_ok, "Falta el parrafo del estado del arte"
    assert concl_ok, "Falta el parrafo de conclusiones"

    # Preservacion: todo heading de E3 sigue en E4
    missing = sorted(e3_headings - e4_headings)
    if missing:
        print("PRESERVACION E3: FALTAN %d encabezados:" % len(missing))
        for m in missing:
            print("   - %s" % m)
        raise AssertionError("No se preservaron todos los encabezados de E3")
    else:
        print("PRESERVACION E3: OK")

    print("== TODAS LAS VERIFICACIONES PASARON ==")
    return {
        "e3_tables": e3_tables,
        "e4_tables": e4_tables,
        "e3_shapes": e3_shapes,
        "e4_shapes": e4_shapes,
    }


def main():
    print("Generando Entrega 4 a partir de la Entrega 3...")
    print("  E3: %s" % E3_PATH)
    print("  E4: %s" % E4_PATH)
    build()
    print("Documento generado y guardado.")
    verify()
    print("Listo: %s" % E4_PATH)


if __name__ == "__main__":
    main()
    sys.exit(0)

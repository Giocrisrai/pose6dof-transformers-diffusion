#!/usr/bin/env python3
"""Reconstruye la sección 'CHARLA divulgativa' del guion, alineada al deck
actual (24 slides) y con tono cercano/técnico que invita a probar y usar.

- Deja INTACTA la sección de defensa.
- Borra desde el Heading 1 'CHARLA divulgativa' hasta el final y reescribe.
- Idempotente: re-ejecutar produce el mismo resultado.

Uso:  ../.venv_thesis/bin/python scripts/build_guion_divulgativa.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[1]
DOC = REPO / "docs/Guion_Completo_Presentaciones.docx"

H1 = "CHARLA divulgativa (Robótica e IA) — ~45 min"

# Cada bloque: (titulo_heading2, [ (kind, texto) ]) con kind in {cue, say, note}
SLIDES = [
    ("Slide 1 — ¿Puede una máquina?", [
        ("cue", "[APERTURA · ~1 min]"),
        ("say", "«Buenas. Les hago una pregunta: ¿cuántos creen que un robot puede "
                "hacer lo que hace un niño de tres años recogiendo juguetes del "
                "suelo?» [deja que respondan, con la mano] «Suena fácil… y es uno de "
                "los problemas más difíciles de la robótica. Hoy van a ver a uno "
                "intentándolo —y lo mejor: todo corre en una laptop, sin "
                "supercomputadora. Al final les dejo cómo probarlo ustedes mismos.»"),
    ]),
    ("Slide 2 — 50 segundos: un robot que ve, piensa y actúa", [
        ("cue", "[EL REEL · ~2 min]"),
        ("say", "«Antes de explicar nada, quiero que lo vean. Cincuenta segundos.»"),
        ("note", "[Reproducir el reel. NO hablar encima. Solo señalar: \"ahí está "
                 "viendo la pieza… ahí decide el camino… ahí la agarra\".]"),
        ("say", "«¿Cómo lo hace? Eso es justo lo que les voy a contar —y van a "
                "entenderlo, prometido.»"),
    ]),
    ("Slide 3 — Lo difícil es fácil", [
        ("cue", "[PARADOJA DE MORAVEC · ~3 min]"),
        ("say", "«Hay una idea preciosa, la paradoja de Moravec: lo que la evolución "
                "tardó millones de años en pulir —ver, mover las manos— es lo más "
                "difícil de copiar en una máquina; mientras que lo \"intelectual\", "
                "como el ajedrez, fue lo primero en caer. Ganarle al campeón mundial: "
                "resuelto en 1997. Recoger una pieza del suelo: todavía es "
                "investigación.»"),
        ("note", "[Margen: \"cualquier bodega del mundo tiene este problema sin "
                 "resolver del todo\".]"),
    ]),
    ("Slide 4 — El reto: bin-picking, el “santo grial” aburrido de la industria", [
        ("cue", "[EL RETO · ~3 min]"),
        ("say", "«El problema se llama bin picking: piezas amontonadas, sin orden, y "
                "el robot tiene que elegir una, agarrarla y soltarla en otro sitio. "
                "Parece simple, pero junta visión 3D, geometría, planificación y "
                "control, todo a la vez. ¿Por qué importa? Porque cada paquete que "
                "reciben de un e-commerce pasó por unas manos haciendo exactamente "
                "esto. Es un mercado de miles de millones.»"),
        ("note", "[La cifra de mercado, dila como \"estimaciones de informes del "
                 "sector\".]"),
    ]),
    ("Slide 5 — Primer problema: ¿dónde está la pieza… y cómo está puesta?", [
        ("cue", "[POSE 6-DoF · ~4 min]"),
        ("say", "«Primer problema: el robot necesita saber no solo DÓNDE está la "
                "pieza —tres números, x, y, z— sino también CÓMO está girada —otros "
                "tres ángulos—. Esos seis números son la \"pose 6-DoF\". Sin ellos, "
                "el robot es una mano que no sabe cómo viene la pieza.»"),
        ("note", "[Interacción: «Tomen un objeto de su mesa. ¿Vieron? Su mano ya "
                 "sabía la orientación antes de tocarlo. Ahora intenten cogerlo con "
                 "el codo bloqueado…» — esa torpeza es un robot sin buena estimación "
                 "de pose.]"),
    ]),
    ("Slide 6 — La tecnología detrás de ChatGPT… mirando en 3D", [
        ("cue", "[FOUNDATIONPOSE · ~4 min]"),
        ("say", "«¿Y cómo calcula esos seis números? Con la misma familia de IA que "
                "hay detrás de ChatGPT: un Transformer. La idea clave es la "
                "\"atención\": en vez de mirarlo todo por igual, el modelo aprende a "
                "fijarse en los detalles que importan, igual que ustedes leen una "
                "palabra en su contexto. Se llama FoundationPose, es de NVIDIA, de "
                "2024. Lo entrenaron con millones de escenas sintéticas, y reconoce "
                "objetos que nunca vio antes —como un chatbot responde preguntas que "
                "nadie le escribió—. Con un sensor de profundidad como el del Face ID "
                "de un móvil le basta para estimar la pose en torno a un segundo. El "
                "panel de la derecha es la telemetría real del sistema.»"),
    ]),
    ("Slide 7 — Ver no basta", [
        ("cue", "[TRANSICIÓN · ~10 s]"),
        ("note", "[Pausa dramática de tres segundos.]"),
        ("say", "«Pero ver no basta. Hay que moverse.»"),
    ]),
    ("Slide 8 — La IA que genera imágenes también genera movimientos", [
        ("cue", "[DIFUSIÓN · ~4 min]"),
        ("say", "«Aquí entra la segunda pieza. ¿Conocen esas IAs que generan "
                "imágenes, como Midjourney? Parten de pura estática, de ruido, y la "
                "van \"limpiando\" hasta que aparece un cuadro. Pues la misma "
                "matemática —los modelos de difusión— sirve para generar movimientos: "
                "en vez de una imagen, lo que va apareciendo del ruido es la "
                "trayectoria del brazo. Mismo truco, otra salida.»"),
    ]),
    ("Slide 9 — No memoriza UNA respuesta: aprende un repertorio", [
        ("cue", "[MULTIMODALIDAD · ~2 min]"),
        ("say", "«Y algo bonito: no memoriza una única respuesta, aprende un "
                "repertorio. Esta nube son cincuenta trayectorias generadas para la "
                "misma pieza, todas distintas y casi todas válidas. Como un humano: "
                "nunca repites exactamente el mismo gesto. Guárdense el detalle: si "
                "puede generar muchas opciones, luego nos quedamos con la mejor.»"),
    ]),
    # ---- NUEVO: glosario ----
    ("Slide 10 — Para entendernos", [
        ("cue", "[GLOSARIO AMABLE · ~2 min · diapositiva nueva]"),
        ("say", "«Hagamos una pausa de un minuto para quedarnos con cinco palabras, "
                "en cristiano. Pose 6-DoF: dónde está algo y cómo está girado. "
                "Transformer: la familia de ChatGPT, que aprende a prestar atención a "
                "lo importante. Modelo de difusión: la familia de los generadores de "
                "imágenes, que parten del ruido y lo pulen. Visual servoing: el robot "
                "se corrige mirando, como tu mano al enhebrar una aguja. Y "
                "bin-picking: sacar piezas amontonadas de una caja.»"),
        ("say", "«Si se quedan con una sola idea, que sea esta: son dos familias de "
                "IA que ustedes ya usan a diario —una para escribir, otra para "
                "imágenes— puestas a trabajar para ver y para mover. Nada de magia.»"),
        ("note", "[Ritmo tranquilo, señalando cada término en pantalla. Es la slide "
                 "que \"da permiso\" a los no técnicos para seguir el resto.]"),
    ]),
    ("Slide 11 — DEMO 1 — Ustedes mandan", [
        ("cue", "[DEMO 1 — GRADIO · ~8 min]"),
        ("say", "«Y ahora mandan ustedes. Díganme dónde está la pieza.»"),
        ("note", "[Ver runbook. Pedir x/y/z a 2-3 personas o usar presets. Comando: "
                 "cd repo_tfm && .venv/bin/python scripts/demo_charla.py → "
                 "http://127.0.0.1:7860. Cada clic genera trayectorias DISTINTAS "
                 "(multimodalidad en vivo).]"),
        ("note", "[MOMENTO WOW: pedir a alguien que pulse el botón rojo \"perturbar "
                 "la pieza\": se desliza y el sistema RE-PLANIFICA en vivo y la agarra "
                 "igual. Frase: \"esto distingue una solución robusta: el entorno "
                 "cambia y el sistema se adapta\".]"),
    ]),
    ("Slide 12 — Todo junto: ojos + cerebro + brazo", [
        ("cue", "[PIPELINE COMPLETO · ~2 min]"),
        ("say", "«Juntemos las piezas. De izquierda a derecha: la cámara ve, el "
                "Transformer estima la pose, la difusión genera la trayectoria y un "
                "control clásico la ejecuta. Ojos, cerebro y brazo. Lo nuevo de este "
                "trabajo es justo la integración: dos familias de IA moderna unidas "
                "por la misma matemática, y publicadas para que cualquiera las use.»"),
    ]),
    ("Slide 13 — DEMO 2 — En vivo en el simulador", [
        ("cue", "[DEMO 2 — COPPELIASIM · ~8 min]"),
        ("say", "«Y ahora, todo el ciclo ejecutándose de verdad en el simulador.»"),
        ("note", "[Ver runbook. Narrar las fases: percepción → planificación (genera "
                 "8 candidatas) → ejecución de la mejor → depósito. Si falla, pasar "
                 "al vídeo de la siguiente slide SIN dramatismo.]"),
    ]),
    # ---- NUEVO: lenguaje natural ----
    ("Slide 14 — Le hablas… y obedece", [
        ("cue", "[LENGUAJE NATURAL · ~2.5 min · diapositiva nueva]"),
        ("say", "«Y aquí viene lo que más me gusta de todo. Hasta ahora le dábamos "
                "coordenadas. Pero, ¿y si simplemente le hablamos? Miren: le digo "
                "\"dame el cubo rojo de la izquierda\"… y lo hace.»"),
        ("note", "[Reproducir el vídeo embebido (39 s, las cinco instrucciones). "
                 "Señalar el rótulo de cada orden y la pieza elegida.]"),
        ("say", "«Lo bonito por dentro: entiende español e inglés, distingue color, "
                "forma, tamaño y hasta el \"de la izquierda\", y si hay varias piezas "
                "elige la correcta. Y lo importante para ustedes: corre en local, sin "
                "mandar nada a la nube, con licencia abierta. Esto es lo que lo "
                "vuelve cercano: no programas… hablas.»"),
    ]),
    ("Slide 15 — El mismo sistema, grabado — con su telemetría real", [
        ("cue", "[VÍDEO CON TELEMETRÍA · ~1 min · plan B o refuerzo]"),
        ("say", "«Y aquí lo tienen grabado, con su telemetría real: ciclos seguidos, "
                "con las latencias y las fases en tiempo real. Sin trucos: lo que "
                "miden los sensores es lo que ven en pantalla.»"),
        ("note", "[Útil también si las demos en vivo funcionaron: muestra el panel de "
                 "métricas con calma.]"),
    ]),
    ("Slide 16 — Lo que nadie cuenta: primero fracasamos", [
        ("cue", "[FRACASO · ~2 min]"),
        ("say", "«Ahora la parte honesta, la que no se suele contar. Mi primer "
                "sistema completo funcionaba un 60 % de las veces. Quise mejorarlo "
                "con aprendizaje por refuerzo… y bajó al 28 %. El robot aprendió a "
                "depositar mejor, pero olvidó cómo agarrar. En IA esto se llama "
                "olvido catastrófico, y le pasa a los mejores laboratorios del "
                "mundo.»"),
    ]),
    ("Slide 17 — La remontada: enseñar como a un aprendiz", [
        ("cue", "[REMONTADA · ~2 min]"),
        ("say", "«Pero la historia no acaba ahí. Diagnostiqué el problema y lo "
                "remonté: un currículo —primero aprender a agarrar, luego a "
                "depositar—, generar ocho opciones y ejecutar la mejor, y arreglar un "
                "error de geometría. Resultado: 84 %, por encima del sistema "
                "original. El mensaje: en IA real se progresa diagnosticando, no por "
                "arte de magia.»"),
    ]),
    ("Slide 18 — ¿Y cuánto costó? Menos que una moto", [
        ("cue", "[COSTE · ~1.5 min]"),
        ("say", "«¿Y cuánto costó todo esto? Menos que una moto. Unos mil novecientos "
                "dólares: una laptop y una GPU de alquiler en la nube, frente a los "
                "quince a ciento cincuenta mil de una celda industrial. La robótica "
                "inteligente se está democratizando: hoy cualquier universidad, pyme o "
                "estudiante motivado puede investigar esto —y, en un rato, "
                "instalárselo.»"),
    ]),
    ("Slide 19 — ¿Qué viene?", [
        ("cue", "[FUTURO · ~1.5 min]"),
        ("say", "«¿Qué viene? Tres cosas: dar el salto del simulador al robot físico; "
                "hacer la difusión diez veces más rápida; y profundizar en los robots "
                "que entienden lenguaje natural, eso que acaban de ver funcionando. "
                "Es la misma receta —percepción más generación de movimiento— que "
                "están usando los robots humanoides de 2025 y 2026.»"),
    ]),
    ("Slide 20 — ¿Qué significa para una Smart Factory? (OCULTA — solo charla de trabajo)", [
        ("cue", "[SOLO CHARLA DE TRABAJO · ~1.5 min]"),
        ("say", "«¿Y qué significa esto para una Smart Factory como la nuestra? Hoy "
                "ustedes miden la planta en tiempo real: OEE, paradas, datos. El "
                "siguiente paso es que la planta además ACTÚE: que un sistema vea, "
                "decida y manipule. Es pasar del dato a la acción física, y de paso "
                "sumar control de calidad por visión. Todo con enfoque de bajo coste "
                "y rigor de software: tests, integración continua, paquetes.»"),
    ]),
    ("Slide 21 — 3 oportunidades concretas (OCULTA — solo charla de trabajo)", [
        ("cue", "[SOLO CHARLA DE TRABAJO · ~1.5 min]"),
        ("say", "«En concreto, veo tres oportunidades para nosotros. Una: picking y "
                "empaque de fin de línea, justo el tipo de clientes que ya tenemos. "
                "Dos: inspección y control de calidad por visión, integrados en el "
                "panel que ya ofrecemos. Y tres: un demostrador de bajo coste como "
                "gancho de I+D y comercial. Con honestidad: queda pendiente el salto a "
                "producción real y la licencia del estimador, pero como prueba de "
                "concepto, la base ya está.»"),
    ]),
    ("Slide 22 — La robótica inteligente ya no es (solo de unos pocos)", [
        ("cue", "[CIERRE · ~30 s]"),
        ("say", "«Quiero cerrar con una idea: la robótica inteligente ya no es "
                "exclusiva de los grandes laboratorios. Todo lo que vieron —código, "
                "datos y vídeos— es público. No se lo lleven solo como una charla "
                "bonita: llévenselo como una herramienta que pueden tocar.»"),
    ]),
    ("Slide 23 — ¿Preguntas?", [
        ("cue", "[Q&A]"),
        ("say", "«¿Preguntas?»"),
        ("note", "[Deja el reel resumen en pantalla. Respuestas típicas en el anexo "
                 "del runbook: sim-to-real, licencia, por qué difusión, coste, "
                 "robustez.]"),
    ]),
    # ---- NUEVO: contacto y uso ----
    ("Slide 24 — Conéctate y úsalo", [
        ("cue", "[LLAMADA A LA ACCIÓN · ~1 min · diapositiva final]"),
        ("say", "«Y me despido con lo más concreto. Si algo de esto les hizo "
                "cosquillas, pruébenlo hoy mismo: una sola línea, \"pip install "
                "bop-bootstrap-ci\", y tienen en su proyecto la parte de métricas que "
                "usé en el TFM. El código completo del pipeline está en GitHub, con "
                "los vídeos y la documentación.»"),
        ("say", "«Tres QR en pantalla: LinkedIn para escribirme, el paquete en PyPI "
                "para instalarlo, y el repositorio para mirarlo por dentro. Si lo "
                "prueban y me cuentan qué tal, me hacen un favor enorme: el mejor "
                "feedback es el que mejora esto y lo hace crecer. Escaneen, "
                "instálenlo… y hablamos. Muchas gracias.»"),
        ("note", "[Deja esta slide en pantalla durante el aplauso y el Q&A: es el "
                 "marco final que quieres que recuerden y fotografíen.]"),
    ]),
]


def _del_from_heading(doc, heading_text):
    """Borra el párrafo del Heading 1 indicado y todos los siguientes."""
    paras = doc.paragraphs
    start = None
    for i, p in enumerate(paras):
        if p.style and p.style.name == "Heading 1" and heading_text in p.text:
            start = i
            break
    if start is None:
        return False
    for p in paras[start:]:
        p._p.getparent().remove(p._p)
    return True


def main() -> int:
    doc = Document(DOC)
    existed = _del_from_heading(doc, "CHARLA divulgativa")
    print(f"  sección divulgativa previa {'eliminada' if existed else 'no encontrada (se añade)'}")

    doc.add_paragraph(H1, style="Heading 1")
    for title, blocks in SLIDES:
        doc.add_paragraph(title, style="Heading 2")
        for kind, text in blocks:
            doc.add_paragraph(text, style="Normal")
    doc.save(DOC)
    print(f"  guion reescrito: {len(SLIDES)} slides divulgativos alineados al deck")
    print(f"  -> {DOC.relative_to(REPO)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
